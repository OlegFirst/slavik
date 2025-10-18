"""
Document Text Extractor
Extracts text and metadata from various document formats

Supported formats:
- PDF (PyMuPDF, pdfplumber, OCR fallback)
- DOCX (python-docx)
- TXT (plain text)
- Excel (openpyxl, pandas)
- Images (pytesseract OCR)

Based on:
- BCM_1/document_processor-обьедененный/processors/extractor.py (lines 18-245)
- BCM_1/document_processor/document_processor.py (lines 87-195)
"""

import os
import re
from typing import Dict, Any, Optional, List
from pathlib import Path
import mimetypes

# PDF Processing
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# DOCX Processing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Excel Processing
try:
    import openpyxl
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# OCR Processing
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


class DocumentExtractor:
    """
    Multi-format document text extractor.

    Extracts:
    - Full text content
    - Metadata (title, author, dates)
    - Page count
    - Word count
    - File properties
    """

    def __init__(self):
        """Initialize extractor with format handlers."""
        self.handlers = {
            'pdf': self._extract_pdf,
            'docx': self._extract_docx,
            'doc': self._extract_docx,
            'txt': self._extract_txt,
            'xlsx': self._extract_excel,
            'xls': self._extract_excel,
            'csv': self._extract_csv,
            'png': self._extract_image,
            'jpg': self._extract_image,
            'jpeg': self._extract_image,
            'tiff': self._extract_image,
        }

    def extract(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and metadata from document.

        Args:
            file_path: Path to document file

        Returns:
            Dictionary containing:
            - text: Extracted full text
            - metadata: Document metadata
            - page_count: Number of pages
            - word_count: Number of words
            - extraction_method: Method used
            - success: Boolean success flag
            - error: Error message if failed

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If format not supported
        """
        if not os.path.exists(file_path):
            return {
                'success': False,
                'error': f'File not found: {file_path}',
                'text': '',
                'metadata': {},
            }

        # Determine file type
        file_ext = Path(file_path).suffix.lower().lstrip('.')

        # Get MIME type
        mime_type, _ = mimetypes.guess_type(file_path)

        # Get handler for file type
        handler = self.handlers.get(file_ext)

        if not handler:
            return {
                'success': False,
                'error': f'Unsupported file format: {file_ext}',
                'text': '',
                'metadata': {},
                'file_type': file_ext,
                'mime_type': mime_type,
            }

        # Execute extraction
        try:
            result = handler(file_path)
            result['success'] = True
            result['file_type'] = file_ext
            result['mime_type'] = mime_type
            result['file_size'] = os.path.getsize(file_path)

            # Calculate word count if not provided
            if 'word_count' not in result and result.get('text'):
                result['word_count'] = len(result['text'].split())

            return result

        except Exception as e:
            return {
                'success': False,
                'error': f'Extraction failed: {str(e)}',
                'text': '',
                'metadata': {},
                'file_type': file_ext,
                'mime_type': mime_type,
            }

    def _extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF using PyMuPDF (primary) or pdfplumber (fallback).

        Based on: BCM_1/document_processor-обьедененный/processors/extractor.py lines 48-112
        """
        result = {
            'text': '',
            'metadata': {},
            'page_count': 0,
            'extraction_method': None,
        }

        # Try PyMuPDF first (faster, better metadata)
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                result['page_count'] = doc.page_count

                # Extract text from all pages
                text_parts = []
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    text_parts.append(page.get_text())

                result['text'] = '\n\n'.join(text_parts)

                # Extract metadata
                metadata = doc.metadata
                result['metadata'] = {
                    'title': metadata.get('title', ''),
                    'author': metadata.get('author', ''),
                    'subject': metadata.get('subject', ''),
                    'keywords': metadata.get('keywords', ''),
                    'creator': metadata.get('creator', ''),
                    'producer': metadata.get('producer', ''),
                    'creation_date': metadata.get('creationDate', ''),
                    'modification_date': metadata.get('modDate', ''),
                }

                result['extraction_method'] = 'pymupdf'
                doc.close()

                return result

            except Exception as e:
                # Fall through to pdfplumber
                pass

        # Try pdfplumber as fallback
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    result['page_count'] = len(pdf.pages)

                    text_parts = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)

                    result['text'] = '\n\n'.join(text_parts)

                    # pdfplumber has limited metadata
                    result['metadata'] = {
                        'title': pdf.metadata.get('Title', ''),
                        'author': pdf.metadata.get('Author', ''),
                    } if pdf.metadata else {}

                    result['extraction_method'] = 'pdfplumber'

                    return result

            except Exception as e:
                raise ValueError(f'PDF extraction failed: {str(e)}')

        raise ValueError('No PDF extraction library available (install PyMuPDF or pdfplumber)')

    def _extract_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from DOCX using python-docx.

        Based on: BCM_1/document_processor-обьедененный/processors/extractor.py lines 114-152
        """
        if not DOCX_AVAILABLE:
            raise ValueError('python-docx not available (install python-docx)')

        result = {
            'text': '',
            'metadata': {},
            'page_count': None,  # DOCX doesn't have fixed pages
            'extraction_method': 'python-docx',
        }

        try:
            doc = DocxDocument(file_path)

            # Extract all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            result['text'] = '\n\n'.join(text_parts)

            # Extract core properties metadata
            core_props = doc.core_properties
            result['metadata'] = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
            }

            # Estimate page count (rough: 500 words per page)
            word_count = len(result['text'].split())
            result['page_count'] = max(1, word_count // 500)

            return result

        except Exception as e:
            raise ValueError(f'DOCX extraction failed: {str(e)}')

    def _extract_txt(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from plain text files.

        Based on: BCM_1/document_processor-обьедененный/processors/extractor.py lines 154-178
        """
        result = {
            'text': '',
            'metadata': {},
            'page_count': 1,
            'extraction_method': 'plain_text',
        }

        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as f:
                result['text'] = f.read()
        except UnicodeDecodeError:
            # Fallback to latin-1
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    result['text'] = f.read()
            except Exception as e:
                raise ValueError(f'Text file reading failed: {str(e)}')

        # Minimal metadata
        result['metadata'] = {
            'encoding': 'utf-8',
        }

        return result

    def _extract_excel(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from Excel files using openpyxl.

        Based on: BCM_1/document_processor-обьедененный/processors/extractor.py lines 180-215
        """
        result = {
            'text': '',
            'metadata': {},
            'page_count': 0,  # Sheets instead of pages
            'extraction_method': 'openpyxl',
        }

        if not OPENPYXL_AVAILABLE:
            raise ValueError('openpyxl not available (install openpyxl)')

        try:
            wb = load_workbook(file_path, read_only=True, data_only=True)
            result['page_count'] = len(wb.sheetnames)

            text_parts = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]

                # Add sheet name as header
                text_parts.append(f'\n=== {sheet_name} ===\n')

                # Extract all cell values
                for row in sheet.iter_rows(values_only=True):
                    row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                    if row_text.strip(' |'):
                        text_parts.append(row_text)

            result['text'] = '\n'.join(text_parts)

            # Excel metadata
            result['metadata'] = {
                'sheets': wb.sheetnames,
                'sheet_count': len(wb.sheetnames),
            }

            wb.close()

            return result

        except Exception as e:
            raise ValueError(f'Excel extraction failed: {str(e)}')

    def _extract_csv(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from CSV files using pandas.

        Based on: BCM_1/document_processor/document_processor.py lines 175-195
        """
        result = {
            'text': '',
            'metadata': {},
            'page_count': 1,
            'extraction_method': 'pandas',
        }

        if not PANDAS_AVAILABLE:
            # Fallback to plain text
            return self._extract_txt(file_path)

        try:
            df = pd.read_csv(file_path)

            # Convert DataFrame to text representation
            result['text'] = df.to_string(index=False)

            # CSV metadata
            result['metadata'] = {
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': df.columns.tolist(),
            }

            return result

        except Exception as e:
            # Fallback to plain text
            return self._extract_txt(file_path)

    def _extract_image(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from images using OCR (pytesseract).

        Based on: BCM_1/document_processor/document_processor.py lines 197-225
        """
        result = {
            'text': '',
            'metadata': {},
            'page_count': 1,
            'extraction_method': 'ocr',
        }

        if not OCR_AVAILABLE:
            raise ValueError('OCR not available (install pytesseract and PIL)')

        try:
            image = Image.open(file_path)

            # Extract text using Tesseract OCR
            result['text'] = pytesseract.image_to_string(image)

            # Image metadata
            result['metadata'] = {
                'format': image.format,
                'mode': image.mode,
                'size': image.size,
                'width': image.width,
                'height': image.height,
            }

            return result

        except Exception as e:
            raise ValueError(f'OCR extraction failed: {str(e)}')


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def clean_extracted_text(text: str) -> str:
    """
    Clean extracted text by removing excess whitespace and artifacts.

    Args:
        text: Raw extracted text

    Returns:
        Cleaned text
    """
    if not text:
        return ''

    # Remove multiple spaces
    text = re.sub(r' +', ' ', text)

    # Remove multiple newlines (keep max 2)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]

    # Remove empty lines at start and end
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()

    return '\n'.join(lines)


def extract_document_sections(text: str) -> List[Dict[str, str]]:
    """
    Extract logical sections from document text.

    Identifies sections based on:
    - Numbered headings (1., 1.1, 1.1.1)
    - ALL CAPS headings
    - Common section keywords

    Args:
        text: Document text

    Returns:
        List of sections with title and content
    """
    sections = []
    current_section = None

    lines = text.split('\n')

    # Patterns for section headers
    numbered_pattern = re.compile(r'^(\d+\.)+\s+(.+)$')
    caps_pattern = re.compile(r'^[A-Z][A-Z\s]{10,}$')

    for line in lines:
        # Check for numbered section
        match = numbered_pattern.match(line.strip())
        if match:
            # Save previous section
            if current_section:
                sections.append(current_section)

            # Start new section
            current_section = {
                'number': match.group(1).rstrip('.'),
                'title': match.group(2).strip(),
                'content': '',
            }
            continue

        # Check for ALL CAPS heading
        if caps_pattern.match(line.strip()) and len(line.strip()) < 100:
            # Save previous section
            if current_section:
                sections.append(current_section)

            # Start new section
            current_section = {
                'number': None,
                'title': line.strip(),
                'content': '',
            }
            continue

        # Add to current section content
        if current_section:
            current_section['content'] += line + '\n'

    # Add final section
    if current_section:
        sections.append(current_section)

    # Clean section content
    for section in sections:
        section['content'] = section['content'].strip()

    return sections
