# -*- coding: utf-8 -*-
"""
Document Processor Microservice for BCM Platform
AI-powered document processing, analysis, and knowledge extraction for BCM documentation
"""
import os
import logging
import asyncio
import hashlib
import mimetypes
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any, Union
from pathlib import Path
import structlog
from fastapi import FastAPI, HTTPException, UploadFile, File, BackgroundTasks, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, Field
import uvicorn
import types

# Optional dependencies wrapped to allow running tests without full stack
try:
    import aiofiles
except Exception:  # pragma: no cover - optional dependency
    aiofiles = None

try:
    import aioredis
except Exception:  # pragma: no cover - optional dependency
    aioredis = types.SimpleNamespace(Redis=object)

try:
    import asyncpg
except Exception:  # pragma: no cover - optional dependency
    asyncpg = None

try:  # pragma: no cover - optional dependency
    from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession
    from sqlalchemy.orm import sessionmaker
except Exception:
    create_async_engine = AsyncSession = sessionmaker = None

# Document processing imports
try:  # pragma: no cover - optional dependency
    import fitz  # PyMuPDF for PDF processing
except Exception:
    fitz = None

try:  # pragma: no cover - optional dependency
    import docx  # python-docx for Word documents
except Exception:
    docx = None

try:  # pragma: no cover - optional dependency
    import openpyxl  # For Excel files
except Exception:
    openpyxl = None

try:  # pragma: no cover - optional dependency
    from PIL import Image
except Exception:
    Image = None

try:  # pragma: no cover - optional dependency
    import pytesseract  # OCR for images
except Exception:
    pytesseract = None

try:  # pragma: no cover - optional dependency
    import textract  # Alternative document extraction
except Exception:
    textract = None

try:  # pragma: no cover - optional dependency
    import spacy  # NLP processing
except Exception:
    spacy = None

try:  # pragma: no cover - optional dependency
    import pandas as pd
except Exception:
    pd = None

# AI/ML imports
try:  # pragma: no cover - optional dependency
    import openai
except Exception:
    openai = None

try:  # pragma: no cover - optional dependency
    from transformers import pipeline, AutoTokenizer, AutoModel
except Exception:
    pipeline = AutoTokenizer = AutoModel = None

try:  # pragma: no cover - optional dependency
    import torch
except Exception:
    torch = None

try:  # pragma: no cover - optional dependency
    import numpy as np
except Exception:
    np = None

try:  # pragma: no cover - optional dependency
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.cluster import KMeans
    from sklearn.metrics.pairwise import cosine_similarity
except Exception:
    TfidfVectorizer = KMeans = cosine_similarity = None

# Configure structured logging
structlog.configure(
    processors=[
        structlog.stdlib.filter_by_level,
        structlog.stdlib.add_logger_name,
        structlog.stdlib.add_log_level,
        structlog.processors.TimeStamper(fmt="iso"),
        structlog.processors.JSONRenderer()
    ],
    context_class=dict,
    logger_factory=structlog.stdlib.LoggerFactory(),
    wrapper_class=structlog.stdlib.BoundLogger,
    cache_logger_on_first_use=True,
)

logger = structlog.get_logger(__name__)

# Pydantic models
class DocumentMetadata(BaseModel):
    filename: str
    file_size: int
    mime_type: str
    file_hash: str
    upload_date: datetime
    processed_date: Optional[datetime] = None
    company_id: str
    document_type: Optional[str] = None
    classification: Optional[str] = None
    language: Optional[str] = None
    page_count: Optional[int] = None

class DocumentContent(BaseModel):
    document_id: str
    raw_text: str
    structured_content: Dict[str, Any]
    extracted_entities: List[Dict[str, Any]]
    key_phrases: List[str]
    summary: Optional[str] = None
    topics: List[str] = []
    compliance_tags: List[str] = []

class BCMDocumentAnalysis(BaseModel):
    document_id: str
    bcm_category: str  # policy, procedure, plan, risk_assessment, bia, exercise
    iso22301_clauses: List[str]
    risk_indicators: List[Dict[str, Any]]
    compliance_score: float
    recommendations: List[str]
    critical_sections: List[Dict[str, Any]]
    stakeholder_references: List[str]
    process_mappings: List[str]

class DocumentSearchRequest(BaseModel):
    query: str
    company_id: str
    document_types: Optional[List[str]] = None
    date_range: Optional[Dict[str, str]] = None
    limit: int = 20

class DocumentComparisonRequest(BaseModel):
    document_id_1: str
    document_id_2: str
    comparison_type: str = "content"  # content, structure, compliance

# Global variables
redis_client: Optional[aioredis.Redis] = None
db_engine = None
nlp_model = None
embeddings_model = None
classification_pipeline = None

# Document processor class
class DocumentProcessor:
    """Advanced document processing for BCM Platform"""
    
    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/msword': self._process_doc,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._process_xlsx,
            'application/vnd.ms-excel': self._process_xls,
            'text/plain': self._process_text,
            'text/csv': self._process_csv,
            'image/png': self._process_image,
            'image/jpeg': self._process_image,
            'image/tiff': self._process_image
        }
        
        # BCM-specific document patterns
        self.bcm_patterns = {
            'policy': ['policy', 'governance', 'framework', 'standard'],
            'procedure': ['procedure', 'process', 'workflow', 'instruction'],
            'plan': ['plan', 'strategy', 'continuity', 'recovery', 'response'],
            'risk_assessment': ['risk', 'assessment', 'analysis', 'threat', 'vulnerability'],
            'bia': ['business impact', 'bia', 'critical', 'dependencies'],
            'exercise': ['exercise', 'drill', 'test', 'simulation', 'tabletop']
        }
        
        # ISO 22301 clause mapping
        self.iso22301_clauses = {
            'context': ['4.1', '4.2', '4.3', '4.4'],
            'leadership': ['5.1', '5.2', '5.3'],
            'planning': ['6.1', '6.2', '6.3'],
            'support': ['7.1', '7.2', '7.3', '7.4', '7.5'],
            'operation': ['8.1', '8.2', '8.3', '8.4'],
            'evaluation': ['9.1', '9.2', '9.3'],
            'improvement': ['10.1', '10.2']
        }
    
    async def process_document(self, file_path: Path, metadata: DocumentMetadata) -> DocumentContent:
        """Process document and extract content"""
        try:
            logger.info(f"Processing document: {metadata.filename}")
            
            # Determine processor based on MIME type
            processor = self.supported_formats.get(metadata.mime_type)
            if not processor:
                raise ValueError(f"Unsupported file format: {metadata.mime_type}")
            
            # Extract text and structure
            extracted_data = await processor(file_path)
            
            # Perform NLP analysis
            nlp_analysis = await self._analyze_text(extracted_data['text'])
            
            # Create document content
            content = DocumentContent(
                document_id=metadata.file_hash,
                raw_text=extracted_data['text'],
                structured_content=extracted_data.get('structure', {}),
                extracted_entities=nlp_analysis['entities'],
                key_phrases=nlp_analysis['key_phrases'],
                summary=nlp_analysis.get('summary'),
                topics=nlp_analysis.get('topics', []),
                compliance_tags=[]
            )
            
            # Perform BCM-specific analysis
            bcm_analysis = await self._analyze_bcm_content(content)
            
            # Store in cache
            if redis_client:
                await redis_client.setex(
                    f"doc:{metadata.file_hash}", 
                    3600, 
                    content.json()
                )
            
            logger.info(f"Document processed successfully: {metadata.filename}")
            return content, bcm_analysis
            
        except Exception as e:
            logger.error(f"Failed to process document {metadata.filename}: {e}")
            raise
    
    async def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF document"""
        try:
            doc = fitz.open(file_path)
            text = ""
            structure = {"pages": [], "metadata": {}}
            
            # Extract metadata
            structure["metadata"] = doc.metadata
            
            # Extract text from each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                text += page_text + "\n"
                
                # Extract page structure
                blocks = page.get_text("dict")["blocks"]
                structure["pages"].append({
                    "page": page_num + 1,
                    "text": page_text,
                    "blocks": len(blocks),
                    "images": len([b for b in blocks if "image" in b])
                })
            
            doc.close()
            return {"text": text, "structure": structure}
            
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            raise
    
    async def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Process Word DOCX document"""
        try:
            doc = docx.Document(file_path)
            text = ""
            structure = {"paragraphs": [], "tables": [], "headers": []}
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
                    structure["paragraphs"].append({
                        "text": para.text,
                        "style": para.style.name if para.style else None
                    })
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                structure["tables"].append(table_data)
                
                # Add table text to main text
                for row in table_data:
                    text += " | ".join(row) + "\n"
            
            # Extract headers/footers
            for section in doc.sections:
                if section.header:
                    header_text = "\n".join([p.text for p in section.header.paragraphs])
                    structure["headers"].append(header_text)
                    text += header_text + "\n"
            
            return {"text": text, "structure": structure}
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            raise
    
    async def _process_doc(self, file_path: Path) -> Dict[str, Any]:
        """Process legacy Word DOC document using textract"""
        try:
            text = textract.process(str(file_path)).decode('utf-8')
            return {"text": text, "structure": {"legacy_format": True}}
        except Exception as e:
            logger.error(f"DOC processing failed: {e}")
            raise
    
    async def _process_xlsx(self, file_path: Path) -> Dict[str, Any]:
        """Process Excel XLSX document"""
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            text = ""
            structure = {"sheets": []}
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_data = []
                
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        row_text = [str(cell) if cell is not None else "" for cell in row]
                        sheet_data.append(row_text)
                        text += " | ".join(row_text) + "\n"
                
                structure["sheets"].append({
                    "name": sheet_name,
                    "rows": len(sheet_data),
                    "data": sheet_data
                })
            
            return {"text": text, "structure": structure}
            
        except Exception as e:
            logger.error(f"XLSX processing failed: {e}")
            raise
    
    async def _process_xls(self, file_path: Path) -> Dict[str, Any]:
        """Process legacy Excel XLS document"""
        try:
            df = pd.read_excel(file_path, sheet_name=None)
            text = ""
            structure = {"sheets": []}
            
            for sheet_name, sheet_df in df.items():
                sheet_text = sheet_df.to_string()
                text += f"Sheet: {sheet_name}\n{sheet_text}\n\n"
                
                structure["sheets"].append({
                    "name": sheet_name,
                    "rows": len(sheet_df),
                    "columns": len(sheet_df.columns),
                    "data": sheet_df.to_dict('records')
                })
            
            return {"text": text, "structure": structure}
            
        except Exception as e:
            logger.error(f"XLS processing failed: {e}")
            raise
    
    async def _process_text(self, file_path: Path) -> Dict[str, Any]:
        """Process plain text document"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                text = await f.read()
            return {"text": text, "structure": {"format": "plain_text"}}
        except Exception as e:
            logger.error(f"Text processing failed: {e}")
            raise
    
    async def _process_csv(self, file_path: Path) -> Dict[str, Any]:
        """Process CSV document"""
        try:
            df = pd.read_csv(file_path)
            text = df.to_string()
            
            structure = {
                "format": "csv",
                "rows": len(df),
                "columns": list(df.columns),
                "data": df.to_dict('records')
            }
            
            return {"text": text, "structure": structure}
            
        except Exception as e:
            logger.error(f"CSV processing failed: {e}")
            raise
    
    async def _process_image(self, file_path: Path) -> Dict[str, Any]:
        """Process image document using OCR"""
        try:
            # Perform OCR
            text = pytesseract.image_to_string(Image.open(file_path))
            
            structure = {
                "format": "image_ocr",
                "confidence": "estimated"  # Could use pytesseract confidence if needed
            }
            
            return {"text": text, "structure": structure}
            
        except Exception as e:
            logger.error(f"Image OCR processing failed: {e}")
            raise
    
    async def _analyze_text(self, text: str) -> Dict[str, Any]:
        """Perform NLP analysis on extracted text"""
        try:
            analysis = {}
            
            # Named Entity Recognition
            if nlp_model:
                doc = nlp_model(text)
                analysis['entities'] = [
                    {
                        "text": ent.text,
                        "label": ent.label_,
                        "start": ent.start_char,
                        "end": ent.end_char
                    } for ent in doc.ents
                ]
            else:
                analysis['entities'] = []
            
            # Key phrase extraction using simple TF-IDF
            try:
                vectorizer = TfidfVectorizer(max_features=20, stop_words='english', ngram_range=(1, 3))
                tfidf_matrix = vectorizer.fit_transform([text])
                feature_names = vectorizer.get_feature_names_out()
                scores = tfidf_matrix.toarray()[0]
                
                # Get top phrases
                phrase_scores = list(zip(feature_names, scores))
                phrase_scores.sort(key=lambda x: x[1], reverse=True)
                analysis['key_phrases'] = [phrase for phrase, score in phrase_scores[:10]]
                
            except Exception as e:
                logger.warning(f"Key phrase extraction failed: {e}")
                analysis['key_phrases'] = []
            
            # Text summarization (if OpenAI is available)
            if len(text) > 1000:  # Only summarize longer texts
                analysis['summary'] = await self._summarize_text(text)
            
            # Topic modeling (simple keyword-based)
            analysis['topics'] = await self._extract_topics(text)
            
            return analysis
            
        except Exception as e:
            logger.error(f"Text analysis failed: {e}")
            return {"entities": [], "key_phrases": [], "topics": []}
    
    async def _summarize_text(self, text: str) -> Optional[str]:
        """Generate text summary using AI"""
        try:
            # Use OpenAI if available
            openai_key = os.getenv('OPENAI_API_KEY')
            if openai_key:
                openai.api_key = openai_key
                
                # Truncate text if too long
                max_tokens = 3000
                if len(text) > max_tokens * 4:  # Rough estimate
                    text = text[:max_tokens * 4]
                
                response = await openai.Completion.acreate(
                    engine="text-davinci-003",
                    prompt=f"Summarize the following BCM document in 3-4 sentences:\n\n{text}",
                    max_tokens=150,
                    temperature=0.3
                )
                
                return response.choices[0].text.strip()
            
            return None
            
        except Exception as e:
            logger.warning(f"Text summarization failed: {e}")
            return None
    
    async def _extract_topics(self, text: str) -> List[str]:
        """Extract topics from text using keyword matching"""
        topics = []
        text_lower = text.lower()
        
        # BCM-related topics
        bcm_topics = {
            'risk_management': ['risk', 'threat', 'vulnerability', 'assessment'],
            'business_continuity': ['continuity', 'recovery', 'resilience', 'restore'],
            'incident_response': ['incident', 'emergency', 'crisis', 'response'],
            'compliance': ['compliance', 'audit', 'iso', '22301', 'standard'],
            'governance': ['governance', 'policy', 'framework', 'oversight'],
            'stakeholders': ['stakeholder', 'communication', 'notification'],
            'resources': ['resource', 'personnel', 'facility', 'technology'],
            'testing': ['test', 'exercise', 'drill', 'validation']
        }
        
        for topic, keywords in bcm_topics.items():
            if any(keyword in text_lower for keyword in keywords):
                topics.append(topic)
        
        return topics
    
    async def _analyze_bcm_content(self, content: DocumentContent) -> BCMDocumentAnalysis:
        """Perform BCM-specific document analysis"""
        try:
            text_lower = content.raw_text.lower()
            
            # Classify document type
            bcm_category = self._classify_bcm_document(text_lower)
            
            # Map to ISO 22301 clauses
            iso_clauses = self._map_iso22301_clauses(text_lower)
            
            # Identify risk indicators
            risk_indicators = self._extract_risk_indicators(content.raw_text)
            
            # Calculate compliance score
            compliance_score = self._calculate_compliance_score(text_lower, iso_clauses)
            
            # Generate recommendations
            recommendations = self._generate_recommendations(bcm_category, compliance_score, risk_indicators)
            
            # Identify critical sections
            critical_sections = self._identify_critical_sections(content)
            
            # Extract stakeholder references
            stakeholders = self._extract_stakeholders(content.extracted_entities)
            
            # Map business processes
            processes = self._map_business_processes(text_lower)
            
            return BCMDocumentAnalysis(
                document_id=content.document_id,
                bcm_category=bcm_category,
                iso22301_clauses=iso_clauses,
                risk_indicators=risk_indicators,
                compliance_score=compliance_score,
                recommendations=recommendations,
                critical_sections=critical_sections,
                stakeholder_references=stakeholders,
                process_mappings=processes
            )
            
        except Exception as e:
            logger.error(f"BCM analysis failed: {e}")
            # Return default analysis
            return BCMDocumentAnalysis(
                document_id=content.document_id,
                bcm_category="unknown",
                iso22301_clauses=[],
                risk_indicators=[],
                compliance_score=0.0,
                recommendations=[],
                critical_sections=[],
                stakeholder_references=[],
                process_mappings=[]
            )
    
    def _classify_bcm_document(self, text: str) -> str:
        """Classify BCM document type"""
        scores = {}
        
        for doc_type, keywords in self.bcm_patterns.items():
            score = sum(text.count(keyword) for keyword in keywords)
            scores[doc_type] = score
        
        if not scores or max(scores.values()) == 0:
            return "unknown"
        
        return max(scores, key=scores.get)
    
    def _map_iso22301_clauses(self, text: str) -> List[str]:
        """Map document content to ISO 22301 clauses"""
        clauses = []
        
        # Keyword mapping to clauses
        clause_keywords = {
            '4.1': ['organization', 'context', 'internal', 'external'],
            '4.2': ['stakeholder', 'interested party', 'requirements'],
            '4.3': ['scope', 'boundary', 'applicability'],
            '4.4': ['bcms', 'management system', 'processes'],
            '5.1': ['leadership', 'commitment', 'top management'],
            '5.2': ['policy', 'objective', 'framework'],
            '5.3': ['roles', 'responsibilities', 'authorities'],
            '6.1': ['risk', 'opportunity', 'assessment'],
            '6.2': ['objective', 'planning', 'achieve'],
            '6.3': ['change', 'planning change'],
            '7.1': ['resource', 'personnel', 'infrastructure'],
            '7.2': ['competence', 'training', 'skill'],
            '7.3': ['awareness', 'communication'],
            '7.4': ['communication', 'internal', 'external'],
            '7.5': ['documented information', 'document'],
            '8.1': ['operation', 'implementation'],
            '8.2': ['business impact', 'bia', 'analysis'],
            '8.3': ['continuity strategy', 'recovery'],
            '8.4': ['exercise', 'test', 'maintain'],
            '9.1': ['monitoring', 'measurement', 'evaluation'],
            '9.2': ['audit', 'internal audit'],
            '9.3': ['management review'],
            '10.1': ['improvement', 'nonconformity'],
            '10.2': ['corrective action', 'continual improvement']
        }
        
        for clause, keywords in clause_keywords.items():
            if any(keyword in text for keyword in keywords):
                clauses.append(clause)
        
        return clauses
    
    def _extract_risk_indicators(self, text: str) -> List[Dict[str, Any]]:
        """Extract risk-related indicators from text"""
        risk_indicators = []
        risk_keywords = [
            'threat', 'vulnerability', 'risk', 'impact', 'likelihood',
            'failure', 'disruption', 'outage', 'breach', 'attack'
        ]
        
        text_lower = text.lower()
        for keyword in risk_keywords:
            count = text_lower.count(keyword)
            if count > 0:
                risk_indicators.append({
                    'indicator': keyword,
                    'frequency': count,
                    'severity': 'medium' if count > 5 else 'low'
                })
        
        return risk_indicators
    
    def _calculate_compliance_score(self, text: str, iso_clauses: List[str]) -> float:
        """Calculate compliance score based on ISO 22301 coverage"""
        total_clauses = len(self.iso22301_clauses)
        covered_clauses = len(set(iso_clauses))
        
        base_score = covered_clauses / total_clauses if total_clauses > 0 else 0
        
        # Bonus for comprehensive coverage
        if 'policy' in text and 'procedure' in text and 'exercise' in text:
            base_score += 0.1
        
        return min(1.0, base_score)
    
    def _generate_recommendations(self, category: str, compliance_score: float, 
                                risk_indicators: List[Dict[str, Any]]) -> List[str]:
        """Generate BCM recommendations"""
        recommendations = []
        
        if compliance_score < 0.5:
            recommendations.append("Consider expanding coverage of ISO 22301 requirements")
        
        if category == 'policy' and compliance_score < 0.7:
            recommendations.append("Policy should reference more BCM framework elements")
        
        if len(risk_indicators) < 3:
            recommendations.append("Consider adding more detailed risk assessment content")
        
        high_risk_count = len([r for r in risk_indicators if r.get('severity') == 'high'])
        if high_risk_count > 5:
            recommendations.append("High concentration of risk indicators - review mitigation strategies")
        
        if category == 'plan' and 'test' not in [r['indicator'] for r in risk_indicators]:
            recommendations.append("Business continuity plans should include testing procedures")
        
        return recommendations
    
    def _identify_critical_sections(self, content: DocumentContent) -> List[Dict[str, Any]]:
        """Identify critical sections in document"""
        critical_sections = []
        
        # Look for sections with high entity density or key BCM terms
        if content.structured_content.get('paragraphs'):
            for i, para in enumerate(content.structured_content['paragraphs']):
                text = para.get('text', '')
                if any(term in text.lower() for term in ['critical', 'essential', 'priority', 'recovery time']):
                    critical_sections.append({
                        'section': f'paragraph_{i}',
                        'content': text[:200],  # First 200 chars
                        'importance': 'high'
                    })
        
        return critical_sections[:5]  # Limit to top 5
    
    def _extract_stakeholders(self, entities: List[Dict[str, Any]]) -> List[str]:
        """Extract stakeholder references from entities"""
        stakeholders = []
        
        # Look for person, organization entities
        for entity in entities:
            if entity.get('label') in ['PERSON', 'ORG', 'ORGANIZATION']:
                stakeholders.append(entity.get('text'))
        
        # Add common BCM stakeholders if mentioned
        text_combined = ' '.join([e.get('text', '') for e in entities])
        bcm_stakeholders = ['management', 'employees', 'customers', 'suppliers', 'regulators']
        
        for stakeholder in bcm_stakeholders:
            if stakeholder in text_combined.lower():
                stakeholders.append(stakeholder)
        
        return list(set(stakeholders))  # Remove duplicates
    
    def _map_business_processes(self, text: str) -> List[str]:
        """Map business processes mentioned in document"""
        processes = []
        
        process_keywords = {
            'it_operations': ['it', 'technology', 'system', 'network'],
            'hr_management': ['human resources', 'personnel', 'staff', 'employee'],
            'finance': ['finance', 'accounting', 'budget', 'payment'],
            'operations': ['operations', 'production', 'manufacturing', 'service'],
            'customer_service': ['customer', 'client', 'support', 'service'],
            'supply_chain': ['supply', 'vendor', 'supplier', 'procurement'],
            'facilities': ['facility', 'building', 'site', 'location'],
            'communications': ['communication', 'media', 'public relations']
        }
        
        for process, keywords in process_keywords.items():
            if any(keyword in text for keyword in keywords):
                processes.append(process)
        
        return processes


# Initialize processor
document_processor = DocumentProcessor()

# Security
security = HTTPBearer()

def verify_api_key(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Verify API key for requests"""
    expected_key = os.getenv('DOCUMENT_PROCESSOR_API_KEY')
    if not expected_key or credentials.credentials != expected_key:
        raise HTTPException(status_code=403, detail="Invalid API key")
    return credentials.credentials

# FastAPI app
app = FastAPI(
    title="BCM Document Processor",
    description="AI-powered document processing and analysis for BCM Platform",
    version="1.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def startup_event():
    """Initialize services on startup"""
    global redis_client, db_engine, nlp_model
    
    logger.info("Initializing Document Processor Service")
    
    # Initialize Redis
    redis_url = os.getenv('REDIS_URL', 'redis://localhost:6379')
    try:
        redis_client = aioredis.from_url(redis_url)
        await redis_client.ping()
        logger.info("Redis connection established")
    except Exception as e:
        logger.warning(f"Redis connection failed: {e}")
    
    # Initialize Database
    db_url = os.getenv('DATABASE_URL', 'postgresql://user:pass@localhost/documents')
    try:
        db_engine = create_async_engine(db_url)
        logger.info("Database connection established")
    except Exception as e:
        logger.warning(f"Database connection failed: {e}")
    
    # Initialize NLP model
    try:
        nlp_model = spacy.load("en_core_web_sm")
        logger.info("SpaCy NLP model loaded")
    except Exception as e:
        logger.warning(f"NLP model loading failed: {e}")
    
    logger.info("Document Processor Service initialized")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "1.0.0",
        "services": {
            "redis": redis_client is not None,
            "database": db_engine is not None,
            "nlp_model": nlp_model is not None
        }
    }

@app.post("/api/v1/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    company_id: str = "",
    document_type: Optional[str] = None,
    background_tasks: BackgroundTasks = BackgroundTasks(),
    api_key: str = Depends(verify_api_key)
):
    """Upload and process document"""
    try:
        # Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        # Create file hash
        content = await file.read()
        file_hash = hashlib.sha256(content).hexdigest()
        
        # Create metadata
        metadata = DocumentMetadata(
            filename=file.filename,
            file_size=len(content),
            mime_type=file.content_type or mimetypes.guess_type(file.filename)[0] or 'application/octet-stream',
            file_hash=file_hash,
            upload_date=datetime.now(),
            company_id=company_id,
            document_type=document_type
        )
        
        # Save file temporarily
        upload_dir = Path(os.getenv('UPLOAD_DIR', '/tmp/documents'))
        upload_dir.mkdir(parents=True, exist_ok=True)
        file_path = upload_dir / f"{file_hash}_{file.filename}"
        
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(content)
        
        # Process document in background
        background_tasks.add_task(
            _process_document_task, 
            file_path, 
            metadata
        )
        
        return {
            "status": "accepted",
            "document_id": file_hash,
            "filename": file.filename,
            "message": "Document processing initiated"
        }
        
    except Exception as e:
        logger.error(f"Document upload failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

async def _process_document_task(file_path: Path, metadata: DocumentMetadata):
    """Background task to process document"""
    try:
        content, bcm_analysis = await document_processor.process_document(file_path, metadata)
        
        # Store results (would typically go to database)
        logger.info(f"Document processed: {metadata.filename}")
        
        # Clean up temporary file
        if file_path.exists():
            file_path.unlink()
            
    except Exception as e:
        logger.error(f"Document processing task failed: {e}")

if __name__ == "__main__":
    port = int(os.getenv('PORT', '8002'))
    host = os.getenv('HOST', '0.0.0.0')
    log_level = os.getenv('LOG_LEVEL', 'info').lower()
    
    uvicorn.run(
        "document_processor:app",
        host=host,
        port=port,
        log_level=log_level,
        reload=False
    )
